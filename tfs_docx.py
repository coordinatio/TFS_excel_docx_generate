from datetime import datetime
from itertools import count
from typing import ItemsView
from tfs import TFSAPI
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm



pat = "icct64e7fzwbpe2fvbo52jq3nsj7c34rdqgp32h7j74ye7ox5jhq" # Ключ доступа

client = TFSAPI("https://tfs.content.ai/", project="HQ/ContentAI",pat=pat) # Соединение с тфс      147233

custom_start = '01-01-2023' # Дата начала
custom_end = '31-01-2023' # Дата конца
tag = 'CRE_12_R6u1' # Тег // query_1 - запрос без тега, query_tags - запрос включает тег

query_1 = """SELECT [System.Title], [Created Date], [Closed Date], [System.AssignedTo], [Tags]
FROM workitems
WHERE [System.State] = 'Done' AND [System.AssignedTo] <> ' '  AND [System.WorkItemType] = 'Task' AND ([Created Date] >= ' """ + custom_start + """ ' AND [Closed Date] <= ' """ + custom_end + """ ')
ORDER BY [System.AssignedTo]
"""

query_tags = """SELECT [System.Title], [Created Date], [Closed Date], [System.AssignedTo], [Tags]
FROM workitems
WHERE [System.State] = 'Done' AND ([Created Date] >= ' """ + custom_start + """ ' AND [Closed Date] <= ' """ + custom_end + """ ') AND [Tags] Contains ' """ + tag + """ '
ORDER BY [System.AssignedTo]
"""

wiql = client.run_wiql(query_1)
workitems = wiql.workitems # Получаем объекты из запроса

members = [] # Массив прикрепленных к задаче людей
tasks_title = [] # Массив описания тасков

dates_start = [] # Массив времени начала таска (default = дата создания таска)
dates_end = [] # Массив времени завершения таска (если есть, иначе None)

for x in workitems: # Выборка данных (Исполнитель, Дата начала/конца, Описание)
    if (x['AssignedTo']):
        members.append(x['AssignedTo'][:x['AssignedTo'].find(' <')]) 
    else:
        members.append(x['AssignedTo'])

    tasks_title.append(x['Title'])   
    dates_start.append(x['CreatedDate'][:10])    
            
    if (x['microsoft.vsts.common.closeddate']):
        dates_end.append(x['microsoft.vsts.common.closeddate'][:10])
    else:
        dates_end.append(x['microsoft.vsts.common.closeddate'])

unique_members = [] # Лист уникальных имен
for member in members:
    if member not in unique_members:
        unique_members.append(member)

for i in range(len(dates_start)): # Перевод даты из str -> datetime
    dates_start[i] = datetime.strptime(dates_start[i], '%Y-%m-%d').date()
    dates_end[i] = datetime.strptime(dates_end[i], '%Y-%m-%d').date()
# ============================Служебное задание==============================================

document = Document() # Создание документа

def make_rows_bold(*rows): # Функция для жирного шрифта в строке 
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

table = document.add_table(1, cols=3, style="Table Grid") # Создание таблицы в 1х3
table.allow_autofit = True # Автоформат текста в ячейке по ширине
head_cells = table.rows[0].cells 

for i, item in enumerate(['Описание', 'Дата начала/конца', 'Исполнитель']): # Заполнение названий столбцов
    p = head_cells[i].paragraphs[0]    
    head_cells[i].text = item
make_rows_bold(table.rows[0])

def getDate(date): # Функция для форматирования даты
    if date:
        return "{0}.{1}.{2}".format(date.day, date.month, date.year)
    else:
        return ""


for i in range(len(unique_members)):    # Заполнение таблицы данными
    row_cells = table.add_row().cells
    min_date = dates_start[0].max
    max_date = dates_end[0].min
    for j in range(len(members)):
        if (members[j] == unique_members[i]):
            if tasks_title[i]:
                row_cells[0].text += tasks_title[j]+";\n"
            else:
                row_cells[0].text = '-'
            if (dates_start[j] < min_date):
                min_date = dates_start[j]
            if (dates_end[j] > max_date):
                max_date = dates_end[j]
    row_cells[1].text = "{0} - {1}".format(getDate(min_date), getDate(max_date))
    row_cells[2].text = unique_members[i]

document.save('Tasks_done.docx') # Сохраняем документ
    




