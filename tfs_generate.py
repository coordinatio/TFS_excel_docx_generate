from itertools import count
from typing import ItemsView
from tfs import TFSAPI
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
import pandas as pd


pat = "icct64e7fzwbpe2fvbo52jq3nsj7c34rdqgp32h7j74ye7ox5jhq" # Ключ доступа

client = TFSAPI("https://tfs.content.ai/", project="HQ/ContentAI",pat=pat) # Соединение с тфс      147233

custom_start = '01-01-2023' # Дата начала
custom_end = '31-01-2023' # Дата конца
tag = 'CRE_12_R6u1' # Тег // query_1 - запрос без тега, query_tags - запрос включает тег

query_1 = """SELECT [System.Title], [Created Date], [Closed Date], [System.AssignedTo], [Tags]
FROM workitems
WHERE [System.State] = 'Done' AND [System.AssignedTo] <> ' '  AND [System.WorkItemType] = 'Task' AND ([Created Date] >= ' """ + custom_start + """ ' AND [Closed Date] <= ' """ + custom_end + """ ')
"""

query_tags = """SELECT [System.Title], [Created Date], [Closed Date], [System.AssignedTo], [Tags]
FROM workitems
WHERE [System.State] = 'Done' AND ([Created Date] >= ' """ + custom_start + """ ' AND [Closed Date] <= ' """ + custom_end + """ ') AND [Tags] Contains ' """ + tag + """ '

"""

wiql = client.run_wiql(query_1)
workitems = wiql.workitems # Получаем объекты из запроса

members = [] # Массив прикрепленных к задаче людей
tasks_title = [] # Массив описания тасков
tags = []

dates_start = [] # Массив времени начала таска (default = дата создания таска)
dates_end = [] # Массив времени завершения таска (если есть, иначе None)

for x in workitems: # Выборка данных (Исполнитель, Дата начала/конца, Описание)
    if (x['AssignedTo']):
        members.append(x['AssignedTo'][:x['AssignedTo'].find(' <')]) 
    else:
        members.append(x['AssignedTo'])

    tasks_title.append(x['Title']) # Title <-    
    dates_start.append(x['CreatedDate'][:10])
    tags.append(x['Tags'])
            
    if (x['microsoft.vsts.common.closeddate']):
        dates_end.append(x['microsoft.vsts.common.closeddate'][:10])
    else:
        dates_end.append(x['microsoft.vsts.common.closeddate'])

# ============================Служебное задание==============================================

document = Document() # Создание документа

def make_rows_bold(*rows): # Функция для жирного шрифта в строке 
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

table = document.add_table(1, cols=3, style="Table Grid") # Создание таблицы в 1х3
table.allow_autofit = True # Автоформат текста в ячейке по ширине
head_cells = table.rows[0].cells 

for i, item in enumerate(['Описание', 'Дата начала/конца', 'Исполнитель']): # Заполнение названий столбцов
    p = head_cells[i].paragraphs[0]    
    head_cells[i].text = item
make_rows_bold(table.rows[0])

def getDate(date): # Функция для форматирования даты
    if date:
        return "{0}-{1}-{2}".format(date[8:10], date[5:7], date[:4])
    else:
        return ""


for i in range(len(members)): # Заполнение таблицы данными
    row_cells = table.add_row().cells
    if tasks_title[i]:
        row_cells[0].text = tasks_title[i]
    else:
        row_cells[0].text = "None"
    row_cells[1].text = "{0} {1}".format(getDate(dates_start[i]), getDate(dates_end[i]))
    if members[i]:
        row_cells[2].text = members[i]
    else:
        row_cells[2].text = "None"    


document.save('test.docx') # Сохраняем документ
    




