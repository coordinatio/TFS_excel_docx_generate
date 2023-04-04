#!/usr/bin/python3
from json import JSONDecoder
from typing import Dict, List, OrderedDict
import zipfile
from tfs import TFSAPI
from docx import Document
import os
import xlsxwriter
import re
import argparse
import datetime
import subprocess
import sys
import unittest


class ArgsTypes:
    @staticmethod
    def validate_names_reference(d: dict):
        msg = 'Names json has incorrect structure, must be flat str->str pairs'
        if type(d) is not dict:
            raise argparse.ArgumentTypeError(msg)
        for k in d:
            if type(k) is not str or type(d[k]) is not str:
                raise argparse.ArgumentTypeError(msg)

    @staticmethod
    def arg_names_reference_file(path_to_file: str) -> dict:
        """reads file, parses json, validates contents
        returns dict with incorrect -> correct name pairs"""
        j = {}
        if not path_to_file:
            return j
        if not os.path.exists(path_to_file):
            raise argparse.ArgumentTypeError(
                "Names reference file %s does not exist" % path_to_file)
        try:
            with open(path_to_file, 'r', encoding="utf-8") as f:
                j = JSONDecoder().decode(f.read())
        except:
            raise argparse.ArgumentTypeError(
                "%s contains invalid json" % path_to_file)
        ArgsTypes.validate_names_reference(j)
        return j

    @staticmethod
    def arg_date(i: str) -> str:
        d = ('-', '.', '/', ' ')
        for x in d:
            try:
                return datetime.datetime.strptime(i, f'%d{x}%m{x}%Y').date().strftime("%d-%m-%Y")
            except:
                pass
        raise argparse.ArgumentTypeError(
            f"Please provide date in either {' or '.join([f'dd{x}mm{x}YYYY' for x in d])} format")


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument('--pat', required=True,
                        help=('A Personal Access Token string which could be obtained '
                              'in the "Security" entry inside your TFS profile'))
    parser.add_argument('--from',
                        type=ArgsTypes.arg_date, metavar='dd-mm-YYYY',
                        help='dd-mm-YYYY', required=True)
    parser.add_argument('--to',
                        type=ArgsTypes.arg_date, metavar='dd-mm-YYYY',
                        help='dd-mm-YYYY', required=True)
    parser.add_argument("--out",
                        default="time_report.xlsx", metavar='./THE_EXCEL_FILE_TO_WRITE_INTO.xlsx',
                        help="File to put the results into. Defaults to 'time_report.xlsx'.")
    parser.add_argument("--open", action='store_true', default=False,
                        help="Tells if to open the resulting file immediately after creation")
    parser.add_argument('--names_reference', type=ArgsTypes.arg_names_reference_file,
                        default='name_filter.json', metavar='./A_SPECIAL_FILE.json',
                        help=("Path to the file containing json with name pairs. "
                              "Defaults to 'name_filter.json'."))
    return parser.parse_args()


class Task:
    def __init__(self, title: str, assignees: List[str], release: str, link: str, date_created=None, date_closed=None) -> None:
        self.title = title
        self.assignees = [x for x in sorted(set(assignees))]
        self.release = release
        self.link = link
        self.date_created = date_created
        self.date_closed = date_closed
        self.broken = not self.title or not self.assignees


class Handler():
    def __init__(self, pat, date_from, date_to) -> None:
        tasks = []
        for w in self.retrieve(pat, date_from, date_to):
            tasks.append(Task(self.get_title(w),
                              self.get_assignees(w),
                              self.get_release(w),
                              self.get_link(w),
                              self.get_date_created(w),
                              self.get_date_closed(w)))
        self.workitems = tasks

    def get_assignees(self, workitem):
        assignees = []
        if (workitem['AssignedTo']):
            assignees.append(
                str(workitem['AssignedTo'][:workitem['AssignedTo'].find(' <')]))
        if workitem['Tags']:
            m = re.search(r'[\@#]([А-Яа-яё]+[_ ][А-Яа-яё]+)', workitem['Tags'])
            if m:
                assignees.append(str(m.group(1).replace('_', ' ')))
        return assignees

    def retrieve(self, pat, date_from, date_to):
        return []

    def get_title(self, workitem):
        return str(workitem['Title'])

    def get_date_created(self, workitem):
        return datetime.datetime.strptime(workitem['CreatedDate'][:10], '%Y-%m-%d').date()

    def get_date_closed(self, workitem):
        return datetime.datetime.strptime(workitem['microsoft.vsts.common.closeddate'][:10], '%Y-%m-%d').date()

    def get_release(self, workitem):
        return ''

    def get_link(self, workitem):
        return str(workitem._links['html']['href'])


class HandlerCai(Handler):
    def retrieve(self, pat, date_from, date_to):
        q1 = f"""SELECT [System.AssignedTo], [Tags]
        FROM workitems
        WHERE
            [System.State] = 'Done'
            AND [System.WorkItemType] = 'Task'
            AND (
                ([Closed Date] >= '{date_from}' AND [Closed Date] <= '{date_to}' AND [Closed Date Override] = '')
                OR
                ([Closed Date Override] >= '{date_from}' AND [Closed Date Override] <= '{date_to}')
                )
            AND [System.Tags] NOT CONTAINS 'EXCLUDE_FROM_TIME_REPORTS'
        ORDER BY [System.AssignedTo]
        """
        w = TFSAPI("https://tfs.content.ai/",
                   project="HQ/ContentAI", pat=pat).run_wiql(q1).workitems

        q2 = f"""SELECT [System.AssignedTo], [Tags]
        FROM workitems
        WHERE
            [System.State] = 'Done'
            AND [System.WorkItemType] = 'Product Backlog Item'
            AND [System.AreaPath] = '%s'
            AND (
                ([Closed Date] >= '{date_from}' AND [Closed Date] <= '{date_to}' AND [Closed Date Override] = '')
                OR
                ([Closed Date Override] >= '{date_from}' AND [Closed Date Override] <= '{date_to}')
                )
            AND [System.Tags] NOT CONTAINS 'EXCLUDE_FROM_TIME_REPORTS'
        ORDER BY [System.AssignedTo]
        """
        for a in ('ContentAI\\Документация', 'ContentAI\\Design'):
            w += TFSAPI("https://tfs.content.ai/",
                        project="HQ/ContentAI",
                        pat=pat).run_wiql(q2 % a).workitems
        return w

    def get_release(self, workitem):
        w = workitem
        while True:
            if w['Tags']:
                m = re.search(r'[A-Z]+_\d+\.\d+\.\d+', w['Tags'])
                if m:
                    return str(m.group(0))
            if not w.parent:
                return ''
            w = w.parent


class HandlerIS(Handler):
    def retrieve(self, pat, date_from, date_to):
        q = f"""SELECT [System.AssignedTo], [Tags]
        FROM workitems
        WHERE
            [System.State] = 'Closed'
            AND ([System.WorkItemType] = 'Bug' OR [System.WorkItemType] = 'Task')
            AND (
                ([Closed Date] >= '{date_from}' AND [Closed Date] <= '{date_to}' AND [Closed Date Override] = '')
                OR
                ([Closed Date Override] >= '{date_from}' AND [Closed Date Override] <= '{date_to}')
                )
            AND [System.Tags] NOT CONTAINS 'EXCLUDE_FROM_TIME_REPORTS'
        ORDER BY [System.AssignedTo]
        """
        return TFSAPI("https://tfs.content.ai/", project="NLC/AIS", pat=pat).run_wiql(q).workitems

    def get_release(self, workitem):
        m = re.search(r'AIS\\(\d+\.\d+)', workitem['system.areapath'])
        if m:
            return "IS_%s" % str(m.group(1))
        return ''


class HandlerLingvo(Handler):
    def retrieve(self, pat, date_from, date_to):
        qs = {'Lingvo':
              f"""SELECT [System.AssignedTo], [Tags]
                FROM workitems
                WHERE
                    [System.State] = 'Closed'
                    AND [System.TeamProject] <> 'lingvo.inbox'
                    AND ([System.WorkItemType] = 'Bug' OR [System.WorkItemType] = 'Feature')
                    AND ([Closed Date] >= '{date_from}' AND [Closed Date] <= '{date_to}')
                    AND [System.Tags] NOT CONTAINS 'EXCLUDE_FROM_TIME_REPORTS'
                ORDER BY [System.AssignedTo]
                """,
              'LingvoLive':
              f"""SELECT [System.AssignedTo], [Tags]
                FROM workitems
                WHERE
                    [System.State] = 'Closed'
                    AND ([System.WorkItemType] = 'Bug' OR [System.WorkItemType] = 'Feature')
                    AND ([Closed Date] >= '{date_from}' AND [Closed Date] <= '{date_to}')
                    AND [System.Tags] NOT CONTAINS 'EXCLUDE_FROM_TIME_REPORTS'
                ORDER BY [System.AssignedTo]
                """}
        w = []
        for p in qs:
            w += TFSAPI("https://tfs.content.ai/", project=p,
                        pat=pat).run_wiql(qs[p]).workitems
        return w

    def get_release(self, workitem):
        spec_a = {'Lingvo X6': 'LX6',
                  'lingvo.mobile.iOS': 'LMI',
                  'lingvo.mobile.android': 'LMA',
                  'lingvo.mac': 'LFM',
                  'lingvo.live.ios': 'LLI'}
        m = re.search(r'(.+?)\\(.+\\)?(\d+\.\d+(\.\d+)?)',
                      workitem['system.iterationpath'])
        if m:
            prj = m.group(1)
            ver = m.group(3)
            if prj in spec_a:
                return '%s_%s' % (spec_a[prj], ver)
        spec_b = {'lingvo.mobile.services': 'LLB',
                  'lingvo.live.services': 'LLB',
                  'lingvo.live.web': 'LLWW'}
        m = re.search(r'(.+?)\\.*', workitem['system.iterationpath'])
        if m:
            prj = m.group(1)
            if prj in spec_b:
                return spec_b[prj]
        return ''


class NameNormalizer:
    def __init__(self, ref: dict) -> None:
        self.dict = ref
        self.vals = {x for x in ref.values()}

    def normalize(self, s: str) -> tuple[str, bool]:
        """returns (normalized name : str, if the name is known : bool)"""
        if s in self.dict:
            return (self.dict[s], True)
        return (s, s in self.vals)


class Matrix:
    class AssigneeInfo:
        def __init__(self, releases_ever_known: set, is_name_known: bool) -> None:
            self.tasks_ttl = 0
            self.releases = OrderedDict()
            for r in releases_ever_known:
                self.releases[r] = []
            self.default = []  # here all not release related tasks go
            self.name_known = is_name_known

        def add_task(self, release: str, task: Task):
            if release:
                self.releases[release].append(task)
            else:
                self.default.append(task)
            self.tasks_ttl += 1

    def __init__(self, tasks: List, names_reference={}):
        self.releases_ever_known = {t.release for t in tasks if t.release}
        nn = NameNormalizer(names_reference)
        self.rows = OrderedDict()
        for t in tasks:
            for a, k in OrderedDict([nn.normalize(x) for x in t.assignees]).items():
                if a not in self.rows:
                    self.rows[a] = Matrix.AssigneeInfo(
                        self.releases_ever_known, k)
                self.rows[a].add_task(t.release, t)
        for x in [y for y in names_reference.values() if y not in self.rows]:
            self.rows[x] = Matrix.AssigneeInfo(self.releases_ever_known, True)


class MatrixPrinter:
    msg_no_tasks = 'Нет задач за отчётный период, исправьте задачи в TFS и перегенерируйте отчёт.'
    msg_person_unknown = 'Имя отсутствует в списках коррекции, сломается автоматизация у бухгалтеров.'

    def print(self, m: Matrix):
        header = [x for x in sorted(m.releases_ever_known)]
        col = 0
        for x in [''] + header + ['DEFAULT']:
            self.brush(col, 0, x)
            col += 1
        row = 0
        for y in m.rows:
            col = 0
            row += 1
            if m.rows[y].tasks_ttl and m.rows[y].name_known:
                self.brush(col, row, y)
            else:
                self.brush_highlight(col, row, y)
                msg = []
                if m.rows[y].tasks_ttl == 0:
                    msg.append(MatrixPrinter.msg_no_tasks)
                if not m.rows[y].name_known:
                    msg.append(MatrixPrinter.msg_person_unknown)
                self.brush_comment(col, row, "\n\n".join(msg))

            for x in header:
                col += 1
                if m.rows[y].tasks_ttl == 0:
                    self.brush_percent(col, row, 0)
                else:
                    self.brush_percent(col, row, len(
                        m.rows[y].releases[x])/m.rows[y].tasks_ttl)
                comment = ["%s: %s\n" % (t.title, t.link)
                           for t in m.rows[y].releases[x]]
                if comment:
                    self.brush_comment(col, row, "\n".join(comment))
            col += 1
            if m.rows[y].tasks_ttl == 0:
                self.brush_percent(col, row, 0)
            else:
                self.brush_percent(col, row, len(
                    m.rows[y].default)/m.rows[y].tasks_ttl)
            comment = ["%s: %s\n" % (t.title, t.link)
                       for t in m.rows[y].default]
            if comment:
                self.brush_comment(col, row, "\n".join(comment))

    def brush(self, col, row, x):
        pass

    def brush_percent(self, col, row, x):
        self.brush(col, row, x)

    def brush_highlight(self, col, row, x):
        self.brush(col, row, x)

    def brush_comment(self, col, row, x):
        pass


class ExcelPrinter(MatrixPrinter):
    def __init__(self, filename: str, date_from: str, date_to: str) -> None:
        self.filename = filename
        self.d_from = date_from
        self.d_to = date_to

    def __enter__(self):
        self.book = xlsxwriter.Workbook(self.filename)
        self.sheet = self.book.add_worksheet(
            f'с {self.d_from} до {self.d_to} вкл.')

        self.fmt_percent = self.book.add_format()
        self.fmt_percent.set_num_format('0.00%')

        self.fmt_gray = self.book.add_format({'font_color': '#eeeeee'})
        self.sheet.conditional_format(0, 0, 999, 999, {'type':     'cell',
                                                       'criteria': '=',
                                                       'value':    0,
                                                       'format':   self.fmt_gray})

        self.fmt_highlight = self.book.add_format({'bg_color': '#ffff7f'})
        return self

    def _helpsheet_write(self):
        s = self.book.add_worksheet('КАК РАБОТАТЬ С ТАБЛИЦЕЙ')
#      indent level, text
        ls = ((0, "Краткая инструкция, как работать с этим отчётом:"),
              (0, "1. Не редактируйте таблицу руками, вместо этого правьте задачи в TFS и перегенерируйте отчёт"),
              (1, "Причина тут в том, что требуется, чтобы отчёты о распределении времени (эта табличка) соответствовали"),
              (1, "служебным заданиям, которые генерируются так же автоматически на основе TFS. Поэтому требуется, чтобы "),
              (1, "TFS, как первоисточник, содержал правильные данные."),
              (0, "2. Если задача попала не в тот выпуск то в TFS это исправляется так:"),
              (1, "+ Для задач в проекте HQ\\ContentAI:"),
              (2, "Надо поставить тег выпуска (напр. FTW_13.3.7) на саму задачу или одного из её родителей вверх по иерархии."),
              (1, "+ Для задач из проектов Lingvo или LingvoLive:"),
              (2, "Надо поместить задачу в итерацию, соответствующую выпуску."),
              (1, "+ Для задач из проекта NLC\\AIS:"),
              (2, "Надо поместить задачу в area, соответствующую выпуску."),
              (0, "3. Если вы обнаружили, что в отчёте лишние задачи (например не относящиеся к сделанной работе)"),
              (1, "то вы можете поставить на них тег EXCLUDE_FROM_TIME_REPORTS и они перестанут попадать в отчёт."),
              (0, "4. Если вы видите, что каких-то задач не хватает. (Опять %username% не перевёл сделанные задачи в Done)"),
              (1, "то вы можете закрыть их и заполнить специальное поле 'Close Date Override', чтобы отнести их к нужному периоду"),
              (1, "в поле нужно записывать дату в формате YYYY-MM-DD"),
              (1, "(поле пока не добавлено в проекты Lingvo, уж больно их много, если поле будет там нужно — напишите Ивану В.)"),
              (0, "5. Если вы видите, что не учтена задача, над которой работало несколько человек (например разработчик-тестировщик"),
              (1, "или несколько разработчиков) то вам нужно на соответствующую задачу добавить тег вида #Имя_Фамилия недостающих авторов и "),
              (1, "перегенерировать отчёт.")
              )
        for i, x in enumerate(ls):
            s.write(i, x[0], x[1])

    def __exit__(self, *args):
        self.sheet.autofit()
        self.sheet.freeze_panes(1, 1)
        self._helpsheet_write()
        self.book.close()

    def brush(self, col, row, x):
        self.sheet.write(row, col, x)

    def brush_percent(self, col, row, x):
        self.sheet.write(row, col, x, self.fmt_percent)

    def brush_highlight(self, col, row, x):
        self.sheet.write(row, col, x, self.fmt_highlight)

    def brush_comment(self, col, row, x):
        self.sheet.write_comment(row, col, x)


class DocxPrinter:

    def make_rows_bold(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    def create_table(self, d: Document) -> Document:
        table = d.add_table(1, cols=3, style="Table Grid")
        table.allow_autofit = True
        head_cells = table.rows[0].cells
        for i, item in enumerate(['Описание', 'Дата начала/конца', 'Исполнитель']):
            p = head_cells[i].paragraphs[0]
            head_cells[i].text = item
        DocxPrinter.make_rows_bold(table.rows[0])
        return table

    def normalize_date(self, a: datetime.date, b: datetime.date):
        return "{0}.{1}.{2} - {3}.{4}.{5}".format(a.day, a.month, a.year, b.day, b.month, b.year)

    def create_zip(self, m: Matrix):
        folder_name = [x for x in sorted(m.releases_ever_known)]
        working_path = os.getcwd()
        if not os.path.exists(working_path+"/TFS_docx"):
            os.mkdir("TFS_docx")
        os.chdir(working_path+"/TFS_docx")
        zippers = working_path+"/TFS_docx"

        for x in folder_name:
            if not os.path.exists(zippers+"/"+x):
                os.mkdir(x)
            os.chdir(zippers+"/"+x)
            for y in m.rows:
                if len(m.rows[y].releases[x]) > 0:
                    docx = Document()
                    table = self.create_table(docx)

                    row_cells = table.add_row().cells
                    min_date = datetime.date.max
                    max_date = datetime.date.min
                    for i in m.rows[y].releases[x]:
                        row_cells[0].text += i.title+";\n"
                        if (i.date_created < min_date):
                            min_date = i.date_created
                        if (i.date_closed > max_date):
                            max_date = i.date_closed
                    row_cells[1].text = self.normalize_date(min_date, max_date)
                    row_cells[2].text = y

                    docx.save("%s.docx" % (y))

            os.chdir(zippers)

        if not os.path.exists(zippers+"/Default"):
            os.mkdir("Default")
        os.chdir(zippers+"/Default")
        for y in m.rows:
            if len(m.rows[y].default) > 0:
                docx = Document()
                table = self.create_table(docx)

                row_cells = table.add_row().cells
                min_date = datetime.date.max
                max_date = datetime.date.min
                for i in m.rows[y].default:
                    row_cells[0].text += i.title+";\n"
                    if (i.date_created < min_date):
                        min_date = i.date_created
                    if (i.date_closed > max_date):
                        max_date = i.date_closed
                row_cells[1].text = self.normalize_date(min_date, max_date)
                row_cells[2].text = y

                docx.save("%s.docx" % (y))
        os.chdir(zippers)

        with zipfile.ZipFile(working_path+"/TFS_zipped.zip", 'w', zipfile.ZIP_DEFLATED) as archive_file:
            for dirpath, dirnames, filenames in os.walk(zippers):
                for filename in filenames:
                    file_path = os.path.join(dirpath, filename)
                    archive_file_path = os.path.relpath(file_path, zippers)
                    archive_file.write(file_path, archive_file_path)


def main():
    a = parse_args()

    i = []
    for x in (HandlerCai, HandlerIS, HandlerLingvo):
        i += x(a.pat, vars(a)["from"], vars(a)["to"]).workitems

    with ExcelPrinter(a.out, vars(a)["from"], vars(a)["to"]) as p:
        p.print(Matrix(i, a.names_reference))

    # test = DocxPrinter()
    # test.create_zip(Matrix(i, a.names_reference))

    if (a.open):
        if sys.platform in ("linux", "linux2"):
            subprocess.call(["xdg-open", os.path.abspath(a.out)])
        else:
            print("--open works only on linux yet")


if __name__ == "__main__":
    main()


class MockWorkitem:
    def __init__(self, d: Dict, link='') -> None:
        self.parent = None
        self.d = d
        self._links = {'html': {'href': link}}

    def __getitem__(self, key):
        return self.d[key]


class TestAssigneeExtraction(unittest.TestCase):
    def test_cai_happyday(self):
        d = {'AssignedTo': 'Алексей Калюжный <CONTENT\\AKalyuzhny>',
             'Tags': '@Федор_Симашев; CC_12.8.0',
             'CreatedDate': '2023-01-01',
             'microsoft.vsts.common.closeddate': '2023-01-01',
             'Title': 'Test title'}

        class X(HandlerCai):
            def retrieve(self, pat, date_from, date_to):
                return [MockWorkitem(d)]
        self.assertEqual(X('', '', '').workitems[0].assignees, [
                         'Алексей Калюжный', 'Федор Симашев'])

    def test_cai_no_tags(self):
        d = {'AssignedTo': 'Алексей Калюжный <CONTENT\\AKalyuzhny>',
             'Tags': None,
             'CreatedDate': '2023-01-01',
             'microsoft.vsts.common.closeddate': '2023-01-01',
             'Title': 'Test title'}

        class X(HandlerCai):
            def retrieve(self, pat, date_from, date_to):
                return [MockWorkitem(d)]
        self.assertEqual(X('', '', '').workitems[0].assignees, [
                         'Алексей Калюжный'])

    def test_cai_hashtag(self):
        d = {'AssignedTo': 'Алексей Калюжный <CONTENT\\AKalyuzhny>',
             'Tags': '#Федор_Симашев; CC_12.8.0',
             'CreatedDate': '2023-01-01',
             'microsoft.vsts.common.closeddate': '2023-01-01',
             'Title': 'Test title'}

        class X(HandlerCai):
            def retrieve(self, pat, date_from, date_to):
                return [MockWorkitem(d)]
        self.assertEqual(X('', '', '').workitems[0].assignees, [
                         'Алексей Калюжный', 'Федор Симашев'])

    def test_assignee_doubling_case(self):
        d = {'AssignedTo': 'Федор Симашев <CONTENT\\Somestring>',
             'Tags': '#Федор_Симашев; CC_12.8.0',
             'Title': 'Test title',
             'CreatedDate': '2023-01-01',
             'microsoft.vsts.common.closeddate': '2023-01-01'}

        class X(HandlerCai):
            def retrieve(self, pat, date_from, date_to):
                return [MockWorkitem(d)]
        self.assertEqual(X('', '', '').workitems[0].assignees, [
                         'Федор Симашев'])


class TestReleaseExtraction(unittest.TestCase):
    def test_cai_happyday(self):
        class X(HandlerCai):
            def retrieve(self, pat, date_from, date_to):
                d = {'AssignedTo': None,
                     'Tags': '@Федор_Симашев; Garbage',
                     'CreatedDate': '2023-01-01',
                     'microsoft.vsts.common.closeddate': '2023-01-01',
                     'Title': 'Test title'}
                w = MockWorkitem(d)
                w.parent = MockWorkitem(
                    {'AssignedTo': '', 'Tags': 'CC_12.8.0',
                     'CreatedDate': '2023-01-01',
                     'microsoft.vsts.common.closeddate': '2023-01-01',
                     'Title': 'T'})
                return [w]
        self.assertEqual('CC_12.8.0', X('', '', '').workitems[0].release)

    def test_cai_no_tags(self):
        class X(HandlerCai):
            def retrieve(self, pat, date_from, date_to):
                d = {'AssignedTo': None,
                     'Tags': None,
                     'CreatedDate': '2023-01-01',
                     'microsoft.vsts.common.closeddate': '2023-01-01',
                     'Title': 'Test title'}
                w = MockWorkitem(d)
                w.parent = MockWorkitem(
                    {'AssignedTo': '', 'Tags': 'CC_12.8.0',
                     'CreatedDate': '2023-01-01',
                     'microsoft.vsts.common.closeddate': '2023-01-01',
                     'Title': 'T'})
                return [w]
        self.assertEqual('CC_12.8.0', X('', '', '').workitems[0].release)

    def test_is_happyday(self):
        class X(HandlerIS):
            def retrieve(self, pat, date_from, date_to):
                d = {'AssignedTo': None,
                     'Tags': '@Федор_Симашев; Garbage',
                     'Title': 'Test title',
                     'CreatedDate': '2023-01-01',
                     'microsoft.vsts.common.closeddate': '2023-01-01',
                     'system.areapath': 'AIS\\5.2'}
                return [MockWorkitem(d)]
        self.assertEqual('IS_5.2', X('', '', '').workitems[0].release)

    def test_lx6_happyday(self):
        class X(HandlerLingvo):
            def retrieve(self, pat, date_from, date_to):
                d = {'AssignedTo': None,
                     'Tags': '@Федор_Симашев; Garbage',
                     'Title': 'Test title',
                     'CreatedDate': '2023-01-01',
                     'microsoft.vsts.common.closeddate': '2023-01-01',
                     'system.iterationpath': 'Lingvo X6\\16.3.1'}
                return [MockWorkitem(d)]
        self.assertEqual('LX6_16.3.1', X('', '', '').workitems[0].release)


class TestMatrix(unittest.TestCase):
    def test_happyday(self):
        t1 = Task('A', ['Petr'], 'FTW_13.3.7', 'http://')
        t2 = Task('B', ['Foma', 'Petr'], 'OMG_13.3.8', 'http://')
        m = Matrix([t1, t2])
        self.assertEqual(m.rows['Petr'].tasks_ttl, 2)
        self.assertEqual(m.rows['Foma'].tasks_ttl, 1)
        self.assertTrue('FTW_13.3.7' in m.rows['Petr'].releases)

    def test_name_normalization(self):
        t1 = Task('A', ['Petr'], 'FTW_13.3.7', 'http://')
        t2 = Task('B', ['Foma', 'Petr'], 'OMG_13.3.8', 'http://')
        t3 = Task('C', ['Ptr'], 'FTW_13.3.7', 'http://')
        m = Matrix([t1, t2, t3], {"Ptr": "Petr", "x": "y"})
        self.assertEqual(m.rows['Petr'].tasks_ttl, 3)
        self.assertEqual(m.rows['Foma'].tasks_ttl, 1)
        self.assertTrue('FTW_13.3.7' in m.rows['Petr'].releases)

    def test_empty_assignee_control(self):
        t1 = Task('A', ['Petr'], 'FTW_13.3.7', 'http://')
        t2 = Task('B', ['Foma', 'Petr'], 'OMG_13.3.8', 'http://')
        m = Matrix([t1, t2], {"Ptr": "Petr", "x": "Empty", "y": "Empty"})
        self.assertEqual(m.rows['Empty'].tasks_ttl, 0)

    def test_same_assignee_several_times(self):
        t1 = Task('A', ['Petr'], 'FTW_13.3.7', 'http://')
        t2 = Task('B', ['Ptr', 'Petr'], 'OMG_13.3.8', 'http://')
        m = Matrix([t1, t2], {"Ptr": "Petr"})
        self.assertEqual(m.rows['Petr'].tasks_ttl, 2)


class TestMatrixPrinter(unittest.TestCase):

    class TestPrinter(MatrixPrinter):
        def __init__(self) -> None:
            self.paper = [['']]
            self.paper_comments = [['']]
            self.paper_highlights = [['']]

        def brush(self, col, row, x):
            TestMatrixPrinter.TestPrinter._print(self.paper, col, row, x)

        def brush_comment(self, col, row, x):
            TestMatrixPrinter.TestPrinter._print(
                self.paper_comments, col, row, x)

        def brush_highlight(self, col, row, x):
            TestMatrixPrinter.TestPrinter._print(
                self.paper_highlights, col, row, x)

        @staticmethod
        def _print(p, col, row, x):
            if len(p) < col + 1:
                p += [[''] for x in range(col + 1 - len(p))]
            for i, l in enumerate(p):  # extend all rows
                if len(l) < row + 1:
                    p[i] += ['' for x in range(row + 1 - len(l))]
            p[col][row] = x

        # nicely print contents into string
        def __str__(self) -> str:
            out = ''
            for row in range(len(self.paper[0])):
                for col in range(len(self.paper)):
                    out += "%s, " % self.paper[col][row]
                out += '\n'
            return out

    def test_sheet(self):
        t1 = Task('A', ['Petr'], 'FTW_13.3.7', 'http://')
        t2 = Task('B', ['Foma', 'Petr'], 'OMG_13.3.8', 'http://')
        l = TestMatrixPrinter.TestPrinter()
        l.print(Matrix([t1, t2], {'P': 'Petr', 'F': 'Foma'}))
        out = [['', 'Petr', 'Foma'],
               ['FTW_13.3.7', 0.5, 0.0],
               ['OMG_13.3.8', 0.5, 1.0],
               ['DEFAULT', 0.0, 0.0]]
        for i, col in enumerate(l.paper):
            self.assertListEqual(out[i], col)

    def test_comments(self):
        t1 = Task('A', ['Petr'], 'FTW_13.3.7', 'http://A')
        t2 = Task('B', ['Foma', 'Petr'], 'OMG_13.3.8', 'http://B')
        l = TestMatrixPrinter.TestPrinter()
        l.print(Matrix([t1, t2], {'x': 'A person with no tasks', 'y': 'Foma'}))
        out = [['', MatrixPrinter.msg_person_unknown, '', MatrixPrinter.msg_no_tasks],
               ['', 'A: http://A\n', '', ''],
               ['', 'B: http://B\n', 'B: http://B\n', '']]
        for i, col in enumerate(l.paper_comments):
            self.assertListEqual(out[i], col)

    def test_hightlights_if_no_tasks(self):
        t1 = Task('A', ['Petr'], 'FTW_13.3.7', 'http://A')
        t2 = Task('B', ['Foma', 'Petr'], 'OMG_13.3.8', 'http://B')
        l = TestMatrixPrinter.TestPrinter()
        l.print(
            Matrix([t1, t2], {'P': 'Petr', 'F': 'Foma', 'x': 'Empty', 'y': 'Empty'}))
        out = [['', '', '', 'Empty']]
        for i, col in enumerate(l.paper_highlights):
            self.assertListEqual(out[i], col)

    def test_hightlights_if_new_name(self):
        t1 = Task('A', ['Petr'], 'FTW_13.3.7', 'http://A')
        t2 = Task('B', ['Foma', 'Petr'], 'OMG_13.3.8', 'http://B')
        l = TestMatrixPrinter.TestPrinter()
        l.print(Matrix([t1, t2], {'F': 'Foma', 'x': 'Empty', 'y': 'Empty'}))
        out = [['', 'Petr', '', 'Empty']]
        for i, col in enumerate(l.paper_highlights):
            self.assertListEqual(out[i], col)

    def test_excel_printer(self):
        t1 = Task('Task 1 with very very long description like you can find in real life',
                  ['Petr'], 'FTW_13.3.7', 'http://task1/adsfakjhdslfkjahdlskfjhaldsfhadf/adlskfj')
        t2 = Task('Task 2', ['Sheph', 'Petr'], 'OMG_13.3.8', 'http://task2')
        t3 = Task('Task 3 with very very long description like you can find in real life',
                  ['Petr'], 'FTW_13.3.7', 'http://task3/asdfupfasdfbdsfdsfasdfadv/asdfefwewdf')
        with ExcelPrinter('test_excel_printer.xlsx', '31-01-2023', '28-02-2023') as printer:
            printer.print(Matrix([t1, t2, t3], {'x': 'Empty'}))


class TestNameFilter(unittest.TestCase):
    def test_filtration(self):
        src = ['a', 'b']
        nf = NameNormalizer({"a": "1", "b": "2", "c": "2"})
        dst = [nf.normalize(x)[0] for x in src]
        self.assertListEqual(['1', '2'], dst)


class TestNamesReferenceValidator(unittest.TestCase):
    def test_valid(self):
        j = '{"a" : "b", "s" : "b", "d" : "f"}'
        ArgsTypes.validate_names_reference(JSONDecoder().decode(j))

    def test_invalid1(self):
        j = '{"a" : "b", "s" : {"x" : "y"}}'
        with self.assertRaises(argparse.ArgumentTypeError):
            ArgsTypes.validate_names_reference(JSONDecoder().decode(j))

    def test_invalid2(self):
        j = '{"a" : "b", "s" : ["x", "y"]}'
        with self.assertRaises(argparse.ArgumentTypeError):
            ArgsTypes.validate_names_reference(JSONDecoder().decode(j))

    def test_invalid3(self):
        j = '["a", "s"]'
        with self.assertRaises(argparse.ArgumentTypeError):
            ArgsTypes.validate_names_reference(JSONDecoder().decode(j))
