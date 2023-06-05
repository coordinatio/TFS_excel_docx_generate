from typing import List, OrderedDict, Dict, Tuple
from math import fsum
from xlsxwriter import Workbook
from docx import Document
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED
from pathlib import Path
from re import match
from datetime import datetime

from src.Task import Task


class NameNormalizer:
    def __init__(self, ref: dict) -> None:
        self.dict = ref
        self.vals = {x for x in ref.values()}

    def normalize(self, s: str) -> tuple[str, bool]:
        """returns (normalized name : str, if the name is known : bool)"""
        if s in self.dict:
            return (self.dict[s], True)
        return (s, s in self.vals)


class Matrix:
    class AssigneeInfo:
        def __init__(self, releases_ever_known: set, is_name_known: bool) -> None:
            self.tasks_ttl = 0
            self.releases = OrderedDict()
            for r in releases_ever_known:
                self.releases[r] = []
            self.default = []  # here all not release related tasks go
            self.name_known = is_name_known

        def add_task(self, release: str, task: Task):
            if release:
                self.releases[release].append(task)
            else:
                self.default.append(task)
            self.tasks_ttl += 1

    def __init__(self, tasks: List[Task], names_reference={}):
        self.releases_ever_known = {t.release for t in tasks if t.release}
        nn = NameNormalizer(names_reference)
        self._rows = OrderedDict()
        for t in tasks:
            for a, k in OrderedDict([nn.normalize(x) for x in t.assignees]).items():
                if a not in self._rows:
                    self._rows[a] = Matrix.AssigneeInfo(
                        self.releases_ever_known, k)
                self._rows[a].add_task(t.release, t)
        for x in [y for y in names_reference.values() if y not in self._rows]:
            self._rows[x] = Matrix.AssigneeInfo(self.releases_ever_known, True)

    def num_tasks_in_release(self, person: str, release: str) -> int:
        if release == 'DEFAULT':
            return len(self._rows[person].default)
        return len(self._rows[person].releases[release])

    def num_tasks_ttl(self, person: str) -> int:
        return self._rows[person].tasks_ttl

    def get_tasks_in_release(self, person: str, release: str) -> list[Task]:
        if release == 'DEFAULT':
            return self._rows[person].default
        return self._rows[person].releases[release]

    def is_assignee_known(self, person: str) -> bool:
        return self._rows[person].name_known

    def list_assignees(self) -> list[str]:
        return [x for x in self._rows]


class MatrixPrinter:
    msg_no_tasks = 'Нет задач за отчётный период, исправьте задачи в TFS и перегенерируйте отчёт.'
    msg_person_unknown = 'Имя отсутствует в списках коррекции, сломается автоматизация у бухгалтеров.'
    msg_control_spends = 'Учтено %.0f%% управленческих затрат времени на выпуск'

    @staticmethod
    def get_release_percents(num_tasks_in_release: int,
                             num_tasks_ttl: int,
                             percents_preallocated_for_release: float,
                             percents_preallocated_ttl: float):
        x = 0.0
        if num_tasks_ttl == 0:
            if percents_preallocated_ttl > 0.0001:
                x = percents_preallocated_for_release / percents_preallocated_ttl
        else:
            x = percents_preallocated_for_release + \
                (num_tasks_in_release / num_tasks_ttl) * \
                (1 - percents_preallocated_ttl)
        return round(x, 7)

    @staticmethod
    def count_releases_of_type(releases_ever_known: set[str], rtype: str) -> int:
        return len([x for x in releases_ever_known if x.startswith(f"{rtype}_")])

    @staticmethod
    def get_release_comment(percents_predefined_for_release: float, tasks: List[Task]) -> str:
        s = '\n'.join(['%s: %s\n' % (t.title, t.link) for t in tasks])
        if percents_predefined_for_release < 0.00001:  # almost zero
            return s
        m = MatrixPrinter.msg_control_spends % (
            percents_predefined_for_release*100)
        if s:
            return f"{m}\n\n{s}"
        return m

    class PredefinedSpend:
        class Metadata:
            def __init__(self, distribution: Dict[str, float], releases_ever_known: set[str]) -> None:
                self._distribution = {}
                for d, p in distribution.items():
                    l = len(
                        [x for x in releases_ever_known if x.startswith(f'{d}_')])
                    self._distribution.update(
                        {r: p/l for r in releases_ever_known if r.startswith(f'{d}_')})
                self._distribution.update(
                    {r: 0.0 for r in releases_ever_known if r not in self._distribution})
                if 'DEFAULT' in distribution:
                    self._distribution['DEFAULT'] = distribution['DEFAULT']
                self.ttl_percent = fsum(
                    [v for k, v in self._distribution.items()])

        def __init__(self, predefined_spend: Dict[str, Dict[str, float]], releases_ever_known: set[str]) -> None:
            self.assignees = {k: MatrixPrinter.PredefinedSpend.Metadata(v, releases_ever_known)
                              for k, v in predefined_spend.items()}

        def get_percents_predefined_for_release(self, person: str, release: str) -> float:
            if person not in self.assignees or release not in self.assignees[person]._distribution:
                return 0.0
            return self.assignees[person]._distribution[release]

        def get_percents_preallocated_ttl(self, person: str) -> float:
            if person not in self.assignees:
                return 0.0
            return self.assignees[person].ttl_percent

    def print(self, m: Matrix, predefined_spend: Dict[str, Dict[str, float]] = {}):
        ps = MatrixPrinter.PredefinedSpend(
            predefined_spend, m.releases_ever_known)
        releases = [x for x in sorted(m.releases_ever_known)]
        col = 0
        # печатаем первую строку, где выпуски
        for x in [''] + releases + ['DEFAULT']:
            self.brush(col, 0, x)
            col += 1
        row = 0
        # идём по строкам
        for person in m.list_assignees():
            col = 0
            row += 1
            # сначала пропечатываем имя человека
            if m.num_tasks_ttl(person) and m.is_assignee_known(person):
                self.brush(col, row, person)
            else:
                self.brush_highlight(col, row, person)
                msg = []
                if m.num_tasks_ttl(person) == 0:
                    msg.append(MatrixPrinter.msg_no_tasks)
                if not m.is_assignee_known(person):
                    msg.append(MatrixPrinter.msg_person_unknown)
                self.brush_comment(col, row, "\n\n".join(msg))
            # теперь идём по выпускам и печатаем, сколько там задач в %
            for release in releases:
                col += 1
                p = MatrixPrinter.get_release_percents(
                    m.num_tasks_in_release(person, release),
                    m.num_tasks_ttl(person),
                    ps.get_percents_predefined_for_release(person, release),
                    ps.get_percents_preallocated_ttl(person))
                if p < 0.0001:
                    self.brush_percent(col, row, 0)
                else:
                    self.brush_percent(col, row, p)
                comment = MatrixPrinter.get_release_comment(
                    ps.get_percents_predefined_for_release(person, release),
                    m.get_tasks_in_release(person, release))
                if comment:
                    self.brush_comment(col, row, comment)
            # печатаем ячейку, соответствующую задачам не попавшим ни в один выпуск
            col += 1
            p = MatrixPrinter.get_release_percents(
                m.num_tasks_in_release(person, 'DEFAULT'),
                m.num_tasks_ttl(person),
                ps.get_percents_predefined_for_release(person, 'DEFAULT'),
                ps.get_percents_preallocated_ttl(person))
            if p < 0.0001:
                self.brush_percent(col, row, 0)
            else:
                self.brush_percent(col, row, p)
            comment = MatrixPrinter.get_release_comment(
                ps.get_percents_predefined_for_release(person, 'DEFAULT'),
                m.get_tasks_in_release(person, 'DEFAULT'))
            if comment:
                self.brush_comment(col, row, comment)

    def brush(self, col, row, x):
        pass

    def brush_percent(self, col, row, x):
        self.brush(col, row, x)

    def brush_highlight(self, col, row, x):
        self.brush(col, row, x)

    def brush_comment(self, col, row, x):
        pass


class ExcelPrinter(MatrixPrinter):
    def __init__(self, output: str | BytesIO, date_from: str, date_to: str) -> None:
        self.output = output
        self.d_from = date_from
        self.d_to = date_to

    def __enter__(self):
        self.book = Workbook(self.output)
        self.sheet = self.book.add_worksheet(
            f'с {self.d_from} до {self.d_to} вкл.')

        self.fmt_percent = self.book.add_format()
        self.fmt_percent.set_num_format('0.00%')

        self.fmt_gray = self.book.add_format({'font_color': '#eeeeee'})
        self.sheet.conditional_format(0, 0, 999, 999, {'type':     'cell',
                                                       'criteria': '=',
                                                       'value':    0,
                                                       'format':   self.fmt_gray})

        self.fmt_highlight = self.book.add_format({'bg_color': '#ffff7f'})
        return self

    def _helpsheet_write(self):
        s = self.book.add_worksheet('КАК РАБОТАТЬ С ТАБЛИЦЕЙ')
#      indent level, text
        ls = ((0, "Краткая инструкция, как работать с этим отчётом:"),
              (0, "1. Не редактируйте таблицу руками, вместо этого правьте задачи в TFS и перегенерируйте отчёт"),
              (1, "Причина тут в том, что требуется, чтобы отчёты о распределении времени (эта табличка) соответствовали"),
              (1, "служебным заданиям, которые генерируются так же автоматически на основе TFS. Поэтому требуется, чтобы "),
              (1, "TFS, как первоисточник, содержал правильные данные."),
              (0, "2. Если задача попала не в тот выпуск то в TFS это исправляется так:"),
              (1, "+ Для задач в проекте HQ\\ContentAI:"),
              (2, "Надо поставить тег выпуска (напр. FTW_13.3.7) на саму задачу или одного из её родителей вверх по иерархии."),
              (1, "+ Для задач из проектов Lingvo или LingvoLive:"),
              (2, "Надо поместить задачу в итерацию, соответствующую выпуску."),
              (1, "+ Для задач из проекта NLC\\AIS:"),
              (2, "Надо поместить задачу в area, соответствующую выпуску."),
              (0, "3. Если вы обнаружили, что в отчёте лишние задачи (например не относящиеся к сделанной работе)"),
              (1, "то вы можете поставить на них тег EXCLUDE_FROM_TIME_REPORTS и они перестанут попадать в отчёт."),
              (0, "4. Если вы видите, что каких-то задач не хватает. (Опять %username% не перевёл сделанные задачи в Done)"),
              (1, "то вы можете закрыть их и заполнить специальное поле 'Close Date Override', чтобы отнести их к нужному периоду"),
              (1, "в поле нужно записывать дату в формате YYYY-MM-DD"),
              (1, "(поле пока не добавлено в проекты Lingvo, уж больно их много, если поле будет там нужно — напишите Ивану В.)"),
              (0, "5. Если вы видите, что не учтена задача, над которой работало несколько человек (например разработчик-тестировщик"),
              (1, "или несколько разработчиков) то вам нужно на соответствующую задачу добавить тег вида #Имя_Фамилия недостающих авторов и "),
              (1, "перегенерировать отчёт.")
              )
        for i, x in enumerate(ls):
            s.write(i, x[0], x[1])

    def __exit__(self, *args):
        self.sheet.autofit()
        self.sheet.freeze_panes(1, 1)
        self._helpsheet_write()
        self.book.close()

    def brush(self, col, row, x):
        self.sheet.write(row, col, x)

    def brush_percent(self, col, row, x):
        self.sheet.write(row, col, x, self.fmt_percent)

    def brush_highlight(self, col, row, x):
        self.sheet.write(row, col, x, self.fmt_highlight)

    def brush_comment(self, col, row, x):
        self.sheet.write_comment(row, col, x)


class ServiceAssignmentsMatrix(Matrix):
    def __init__(self, tasks: List[Task], names_reference={}):
        super().__init__(tasks, names_reference)
        self._releases: dict[str, dict[str, List[Tuple[str, str]]]] = dict()
        for r in self.releases_ever_known | {'DEFAULT'}:
            for a in self.list_assignees():
                tasks = self.get_tasks_in_release(a, r)
                if tasks:
                    if r not in self._releases:
                        self._releases[r] = dict()
                    self._releases[r][a] = [
                        (t.essence, t.essence_completed) for t in tasks]

    def list_releases(self) -> List[str]:
        return [k for k in self._releases]

    def list_assignees_by_release(self, release: str) -> List[str]:
        return [k for k in self._releases[release]]

    def list_essences(self, release: str, assignee: str) -> List[str]:
        return [x[0] for x in self._releases[release][assignee]]

    def list_completed_essences(self, release: str, assignee: str) -> List[str]:
        return [x[1] for x in self._releases[release][assignee]]


def get_product_from_release(release: str) -> str:
    m = match(r'([A-Z\d]+)(_\d+\.\d+\.\d+)?', release)
    if m:
        return m.group(1)
    raise RuntimeError(
        f"Unable to extract product from relese '{release}'")


def docx_move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def docx_delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def docx_form_table(docx, assignee: str, date_from: str, date_to: str, tasks: List[str]):
    table = docx.add_table(1, cols=3, style="Table Grid")
    table.allow_autofit = True
    head_cells = table.rows[0].cells
    for i, item in enumerate(['Описание', 'Дата начала/конца', 'Исполнитель']):
        head_cells[i].text = item
    row_cells = table.add_row().cells
    row_cells[0].text = ';\n\n'.join(tasks)
    row_cells[1].text = f"{date_from} - {date_to}"
    row_cells[2].text = assignee
    return table


class DocsGenerator:
    def __init__(self, dir_templates: str) -> None:
        self.dir_templates = Path(dir_templates)

    def locate_template(self, doctype: str, release: str) -> Path:
        if doctype not in ('todo', 'done'):
            raise ValueError(("The doctype must be either 'todo' or 'done',"
                              f" but got '{doctype}'"))
        p = get_product_from_release(release)
        o = self.dir_templates / doctype / f'{p}.docx'
        if o.exists() and o.is_file():
            return o
        raise RuntimeError(f"Template {o.absolute()} is not found")

    def is_template_valid(self, template: Path) -> bool:
        docx = Document(template)
        must = {'%INSERT_THE_TABLE_HERE%': False, '%ASSIGNEE%': False}
        for p in docx.paragraphs:
            for k in must:
                if p.text.find(k) != -1:
                    must[k] = True
        for v in must.values():
            if not v:
                return False
        return True

    def get_docx(self, doctype: str, release: str, assignee: str,
                 date_from: str, date_to: str,
                 tasks: List[str]):

        template = self.locate_template(doctype, release)
        if not self.is_template_valid(template):
            raise RuntimeError(f'Invalid template: {template}')
        docx = Document(template)

        table = docx_form_table(docx, assignee, date_from, date_to, tasks)
        for p in docx.paragraphs:
            if p.text.find('%INSERT_THE_TABLE_HERE%') != -1:
                docx_move_table_after(table, p)
                docx_delete_paragraph(p)
                break

        replace_map = {'%ASSIGNEE%': assignee,
                       '%DATE_FROM%': date_from}
        for p in docx.paragraphs:
            for k, v in replace_map.items():
                if p.text.find(k) != -1:
                    p.text = p.text.replace(k, v)
        return docx


def get_bundle_zip(sam: ServiceAssignmentsMatrix, date_from: str, date_to: str, predef_spend, dg: DocsGenerator) -> bytes:
    z = BytesIO()
    with ZipFile(z, 'a', ZIP_DEFLATED, False) as zf:
        o = BytesIO()
        with ExcelPrinter(o, date_from, date_to) as p:
            p.print(sam, predef_spend)
        zf.writestr(f'{date_from}-{date_to}.xlsx', o.getvalue())
        for r in sam.list_releases():
            for a in sam.list_assignees_by_release(r):
                o = BytesIO()
                dg.get_docx('todo', r, a, date_from, date_to,
                            sam.list_essences(r, a)).save(o)
                zf.writestr(f'{r}/todo/{a}.docx', o.getvalue())
                o = BytesIO()
                dg.get_docx('done', r, a, date_from, date_to,
                            sam.list_completed_essences(r, a)).save(o)
                zf.writestr(f'{r}/done/{a}.docx', o.getvalue())
    return z.getvalue()
